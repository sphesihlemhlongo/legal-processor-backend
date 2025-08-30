import os
from typing import Optional
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None

from logger_config import get_logger

logger = get_logger(__name__)

class DocumentWriter:
    """Handles writing processed content to various output formats"""
    
    def __init__(self):
        self.output_dir = "outputs"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def write_txt(self, content: str, filename: str) -> str:
        """
        Write content to a text file
        
        Args:
            content: Text content to write
            filename: Output filename
            
        Returns:
            Path to written file
        """
        try:
            file_path = os.path.join(self.output_dir, filename)
            
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(content)
            
            logger.info(f"Text file written: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error writing text file {filename}: {str(e)}")
            raise
    
    def write_docx(self, content: str, filename: str) -> str:
        """
        Write content to a DOCX file with proper formatting
        
        Args:
            content: Text content to write
            filename: Output filename
            
        Returns:
            Path to written file
        """
        try:
            if not Document:
                raise ImportError("python-docx not available")
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Processed Legal Document', 0)
            
            # Process content by sections
            sections = content.split('\n\n')
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                # Check if it's a heading (starts with # or ##)
                if section.startswith('#'):
                    heading_text = section.lstrip('#').strip()
                    level = min(section.count('#'), 3)  # Max heading level 3
                    doc.add_heading(heading_text, level)
                
                # Check if it's a bullet point section
                elif '•' in section or section.startswith('-'):
                    lines = section.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line.startswith('•') or line.startswith('-'):
                            # Add as bullet point
                            p = doc.add_paragraph()
                            p.style = 'List Bullet'
                            p.add_run(line.lstrip('•-').strip())
                        elif line:
                            # Regular paragraph
                            doc.add_paragraph(line)
                else:
                    # Regular paragraph
                    doc.add_paragraph(section)
            
            # Save document
            doc.save(file_path)
            
            logger.info(f"DOCX file written: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error writing DOCX file {filename}: {str(e)}")
            raise
    
    def write_pdf(self, content: str, filename: str) -> str:
        """
        Write content to a PDF file (optional enhancement)
        
        Note: This is a placeholder for PDF generation
        For production, consider using libraries like reportlab or weasyprint
        """
        try:
            # For now, write as text and note PDF limitation
            txt_filename = filename.replace('.pdf', '.txt')
            file_path = self.write_txt(content, txt_filename)
            
            logger.info(f"PDF generation not implemented, saved as TXT: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error writing PDF file {filename}: {str(e)}")
            raise
    
    def create_filename(self, original_name: str, output_type: str, extension: str = "docx") -> str:
        """
        Create standardized filename for outputs
        
        Args:
            original_name: Original document name
            output_type: 'plainEnglish' or 'summary'
            extension: File extension
            
        Returns:
            Formatted filename
        """
        # Remove original extension
        base_name = Path(original_name).stem
        
        # Clean filename for filesystem compatibility
        base_name = re.sub(r'[^\w\-_\.]', '_', base_name)
        
        return f"{base_name}_{output_type}.{extension}"
    
    def cleanup_temp_files(self, file_paths: list) -> None:
        """Remove temporary files after processing"""
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.debug(f"Cleaned up temp file: {file_path}")
            except Exception as e:
                logger.warning(f"Could not remove temp file {file_path}: {str(e)}")