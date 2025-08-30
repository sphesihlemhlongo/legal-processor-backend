import os
import logging
from typing import Dict, Any
from pathlib import Path

try:
    import pdfplumber
    import PyPDF2
    import fitz  # PyMuPDF
except ImportError:
    pdfplumber = None
    PyPDF2 = None
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

from logger_config import get_logger

logger = get_logger(__name__)

class DocumentReader:
    """Handles reading various document formats and extracting structured text"""
    
    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".txt"]
        
    def read_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Read a document and return structured content
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename for format detection
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            file_extension = Path(filename).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            logger.info(f"Reading document: {filename} ({file_extension})")
            
            if file_extension == ".pdf":
                return self._read_pdf(file_path)
            elif file_extension == ".docx":
                return self._read_docx(file_path)
            elif file_extension == ".txt":
                return self._read_txt(file_path)
            else:
                raise ValueError(f"No reader available for {file_extension}")
                
        except Exception as e:
            logger.error(f"Error reading document {filename}: {str(e)}")
            raise
    
    def _read_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using multiple libraries for best results"""
        content = {"text": "", "sections": [], "metadata": {}}
        
        try:
            # Try pdfplumber first (best for structured text)
            if pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                    
                    content["text"] = "\n\n".join(text_parts)
                    content["metadata"]["pages"] = len(pdf.pages)
                    
            # Fallback to PyMuPDF if pdfplumber fails
            elif fitz:
                doc = fitz.open(file_path)
                text_parts = []
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                
                content["text"] = "\n\n".join(text_parts)
                content["metadata"]["pages"] = len(doc)
                doc.close()
                
            # Final fallback to PyPDF2
            elif PyPDF2:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                    
                    content["text"] = "\n\n".join(text_parts)
                    content["metadata"]["pages"] = len(pdf_reader.pages)
            else:
                raise ImportError("No PDF reading library available")
                
        except Exception as e:
            logger.error(f"PDF reading error: {str(e)}")
            raise
        
        return content
    
    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX maintaining structure"""
        content = {"text": "", "sections": [], "metadata": {}}
        
        try:
            if not Document:
                raise ImportError("python-docx not available")
                
            doc = Document(file_path)
            text_parts = []
            sections = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if paragraph is a heading
                    if paragraph.style.name.startswith('Heading'):
                        sections.append({
                            "type": "heading",
                            "level": paragraph.style.name,
                            "text": paragraph.text.strip()
                        })
                        text_parts.append(f"\n## {paragraph.text.strip()}\n")
                    else:
                        text_parts.append(paragraph.text.strip())
            
            content["text"] = "\n\n".join(text_parts)
            content["sections"] = sections
            content["metadata"]["paragraphs"] = len(doc.paragraphs)
            
        except Exception as e:
            logger.error(f"DOCX reading error: {str(e)}")
            raise
            
        return content
    
    def _read_txt(self, file_path: str) -> Dict[str, Any]:
        """Read plain text file"""
        content = {"text": "", "sections": [], "metadata": {}}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                content["text"] = text
                content["metadata"]["characters"] = len(text)
                content["metadata"]["lines"] = len(text.splitlines())
                
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                    content["text"] = text
                    content["metadata"]["characters"] = len(text)
                    content["metadata"]["lines"] = len(text.splitlines())
                    content["metadata"]["encoding"] = "latin-1"
            except Exception as e:
                logger.error(f"Text reading error: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Text reading error: {str(e)}")
            raise
            
        return content